#!/usr/bin/env python3
"""Render buyer-email.md and owner-email.md as Word .docx files.

Why: markdown bold/italic don't render in mail clients, and a .docx
deliverable is what we hand off. The .md remains the source of truth;
this script generates the matching .docx.

Supports the markdown features we actually use:
  - # / ## headings
  - blank-line paragraph breaks
  - **bold** and *italic* inline
  - [link text](https://...) inline hyperlinks
  - - bulleted list items (one per line)
  - --- horizontal rule
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path("/home/andreahimmel/lrp-tx-gis")
SRC_DIR = REPO / "outputs" / "reports"
SRCS = [
    (SRC_DIR / "buyer-email.md",  SRC_DIR / "buyer-email.docx",
        "CONFIDENTIAL — CARAMBA NORTH · DRAFT FOR REVIEW"),
    (SRC_DIR / "owner-email.md",  SRC_DIR / "owner-email.docx",
        "CONFIDENTIAL — CARAMBA NORTH · INTERNAL DRAFT"),
]

INLINE_RE = re.compile(
    r"(\*\*[^*]+?\*\*"     # **bold**
    r"|\*[^*\n]+?\*"        # *italic* (single line)
    r"|\[[^\]]+?\]\([^)]+?\))"  # [text](url)
)

def add_hyperlink(paragraph, url, text, size_pt=11, color=(0x05, 0x63, 0xC1)):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "%02X%02X%02X" % color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size_pt * 2)); rPr.append(sz)
    rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), "Calibri"); rFonts.set(qn("w:hAnsi"), "Calibri"); rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; r.append(t)
    h.append(r)
    paragraph._p.append(h)

def add_inline(p, text, size=11, italic_default=False):
    """Parse markdown inline (bold/italic/links) into runs on paragraph p."""
    tokens = INLINE_RE.split(text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) >= 4:
            r = p.add_run(tok[2:-2]); r.bold = True
            r.font.name = "Calibri"; r.font.size = Pt(size)
            if italic_default: r.italic = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) >= 2 and not tok.startswith("**"):
            r = p.add_run(tok[1:-1]); r.italic = True
            r.font.name = "Calibri"; r.font.size = Pt(size)
        elif tok.startswith("[") and "](" in tok and tok.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            if m:
                add_hyperlink(p, m.group(2), m.group(1), size_pt=size)
                continue
            r = p.add_run(tok)
            r.font.name = "Calibri"; r.font.size = Pt(size)
        else:
            r = p.add_run(tok)
            r.font.name = "Calibri"; r.font.size = Pt(size)
            if italic_default: r.italic = True

def new_para(doc, *, italic=False, size=11, space_after=8, line_spacing=1.2, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def horizontal_rule(doc):
    """Insert a thin grey horizontal-rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "B0B0B0")
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.space_before = Pt(6)

def render(md_text, doc):
    # Split into blocks separated by one-or-more blank lines.
    blocks = re.split(r"\n\s*\n", md_text.strip())
    for block in blocks:
        block = block.rstrip()
        if not block:
            continue
        # H1
        if block.startswith("# "):
            doc.add_heading(block[2:].strip(), level=0)
            continue
        if block.startswith("## "):
            doc.add_heading(block[3:].strip(), level=1)
            continue
        # Horizontal rule
        if block.strip() == "---":
            horizontal_rule(doc)
            continue
        # Italic-Notes block: starts with *Notes (..., ends with *
        if block.startswith("*") and block.endswith("*") and not block.startswith("**"):
            inner = block[1:-1]
            p = new_para(doc, italic=True, size=9, space_after=4)
            add_inline(p, inner, size=9, italic_default=True)
            continue
        # Bullet list block (every line starts with "- ")
        lines = block.split("\n")
        if all(line.lstrip().startswith("- ") for line in lines):
            for line in lines:
                txt = line.lstrip()[2:]
                p = new_para(doc, style="List Bullet", space_after=4)
                add_inline(p, txt)
            continue
        # Email-header block: lines like "**Field:** value"
        header_field_re = re.compile(r"^\*\*([A-Z][^:]+):\*\* (.+)$")
        if all(header_field_re.match(line) for line in lines):
            for line in lines:
                m = header_field_re.match(line)
                p = new_para(doc, space_after=2)
                r = p.add_run(f"{m.group(1)}: ")
                r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(11)
                add_inline(p, m.group(2))
            continue
        # Default: regular paragraph (may contain inline bold / italic / links)
        # If multi-line block, keep as single paragraph (markdown convention: joins with space)
        combined = " ".join(line.strip() for line in lines)
        p = new_para(doc)
        add_inline(p, combined)

def setup_doc(classification_text):
    doc = Document()
    s = doc.styles["Normal"]
    s.font.name = "Calibri"; s.font.size = Pt(11)
    s.paragraph_format.line_spacing = 1.2
    for hn in ("Heading 1", "Heading 2", "Title", "List Bullet"):
        try:
            st = doc.styles[hn]
            if st.font is not None:
                st.font.name = "Calibri"
        except KeyError:
            pass
    # Running header: classification stripe in red.
    hdr = doc.sections[0].header
    hp = hdr.paragraphs[0]
    hr = hp.add_run(classification_text)
    hr.bold = True; hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0xB0, 0, 0)
    return doc

def main():
    for src, out, classification in SRCS:
        md_text = src.read_text(encoding="utf-8")
        doc = setup_doc(classification)
        render(md_text, doc)
        doc.save(out)
        print(f"wrote {out}  ({out.stat().st_size:,} bytes)")

if __name__ == "__main__":
    main()
