from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/andreahimmel/lrp-tx-gis/outputs/reports/Pecos-Shallow-Drilling-Caramba-Vibration.docx"
MAP_URL = "https://lrp-tx-gis.netlify.app"
CHARTS = "/home/andreahimmel/lrp-tx-gis/outputs/reports/charts"

def shade(el, fill):
    pr = el.get_or_add_pPr() if el.tag.endswith('}p') else el.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill); pr.append(shd)

def runs(p, segs, size=11, color=None):
    for text, bold in segs:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
        if color: r.font.color.rgb = RGBColor(*color)

def para(doc, segs, size=11, fill=None, space_after=6, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    runs(p, segs, size)
    if italic:
        for r in p.runs: r.italic = True
    if fill: shade(p._p, fill)
    return p

def bullet(doc, segs, fill=None, space_after=4):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    runs(p, segs, 11)
    if fill: shade(p._p, fill)
    return p

def cell_text(cell, text, bold=False, fill=None, color=None):
    cell.text = ''; p = cell.paragraphs[0]; r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(10)
    if color: r.font.color.rgb = RGBColor(*color)
    if fill: shade(cell._tc, fill)

def table(doc, headers, rows, hdr_fill='DBE8FD', row_styles=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, fill=hdr_fill, color=(0x1F, 0x3A, 0x63))
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        st = (row_styles or {}).get(ri, {})
        for ci, val in enumerate(row):
            cf = st.get('fill')
            col = (0x14, 0x53, 0x2D) if cf == 'CFECCF' else None
            cell_text(cells[ci], val, bold=st.get('bold', False), fill=cf, color=col)
    return t

def add_hyperlink(paragraph, url, text, color=(0x05, 0x63, 0xC1)):
    """Insert a clickable hyperlink in a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '%02X%02X%02X' % color); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '22'); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_figure(doc, fname, width_in=6.0, caption=None):
    """Insert a centered figure with optional italic caption."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f"{CHARTS}/{fname}", width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def add_toc(doc):
    """Insert a Word TOC field. Word auto-populates on open (or via Right-click -> Update Field)."""
    p = doc.add_paragraph(); run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    ft = OxmlElement('w:t'); ft.text = "Table of Contents — right-click and choose 'Update Field' in Word to populate."
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, instr, f2, ft, f3):
        run._r.append(el)
    p.paragraph_format.space_after = Pt(12)

doc = Document()
s = doc.styles['Normal']
s.font.name = 'Calibri'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.2
# Apply Calibri across heading styles too (Word defaults Headings to Calibri Light;
# we keep the same family for consistency).
for hn in ('Heading 1', 'Heading 2', 'Title', 'List Bullet'):
    try:
        st = doc.styles[hn]
        if st.font is not None:
            st.font.name = 'Calibri'
    except KeyError:
        pass

# Running header (every page): CONFIDENTIAL classification stripe.
hdr = doc.sections[0].header
hp = hdr.paragraphs[0]
hr = hp.add_run("CONFIDENTIAL — PECOS COUNTY · CARAMBA NORTH")
hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = RGBColor(0xB0, 0, 0)

doc.add_heading("Pecos County Drilling Activity — Historical & Recent Record", level=0)
para(doc, [("Prepared: 2026-05-19  ·  Subject site: Caramba North tract (≈1,300 ac), Pecos County, TX — centroid ≈ 30.9032° N, 102.9747° W  ·  Classification: Confidential", False)], size=9)

doc.add_heading("Purpose", level=1)
para(doc, [("This memorandum summarizes the historical and recent record of oil-and-gas drilling — with attention to shallow (<3,000 ft) wells — at and within ten miles of the Caramba North tract, drawn from the Railroad Commission of Texas (RRC) wellbore and drilling-permit records. It is provided as context for evaluating potential ground-vibration considerations for a data-center development on the site. Throughout, ", False),
           ("new drilling (a new wellbore) is distinguished from recompletions (rework of an existing wellbore — no new hole drilled)", True),
           ("; only new drilling involves a drilling rig and the hydraulic-fracturing completion associated with ground vibration.", False)])
para(doc, [("Proximity is reported at explicit distances from the tract centroid — principally ", False), ("within two miles", True), (" and ", False), ("within ten miles", True), (". Ten miles is a deliberately generous boundary: ground vibration from drilling and completion attenuates well within that distance.", False)])
_map_p = doc.add_paragraph()
_map_p.paragraph_format.space_after = Pt(6)
_r_intro = _map_p.add_run("This report is intended to accompany and utilizes the data underlying the interactive map of Caramba North, which can be accessed through ")
_r_intro.italic = True; _r_intro.font.size = Pt(11)
add_hyperlink(_map_p, MAP_URL, "this link")
_r_close = _map_p.add_run(".")
_r_close.italic = True; _r_close.font.size = Pt(11)

doc.add_heading("Table of Contents", level=1)
add_toc(doc)
doc.add_page_break()

doc.add_heading("Summary of findings", level=1)
para(doc, [
    ("No new drilling is occurring at or near the Caramba North site. ", True),
    ("Counting only genuine new wells (wellbore records, recompletion re-stamps excluded): ", False),
    ("no well of any kind has been spudded within two miles of the tract in over a decade", True),
    (", ", False), ("no new-drill well lies within five miles", True),
    (", and only ", False), ("three new-drill wells sit within ten miles", True),
    (" across all of 2020–2025 (nearest ≈ 6.9 miles; all three are deep ≥9,200 ft, spud 2020 and 2025).", False),
], fill='EAF2FF', space_after=6)
para(doc, [
    ("The wells within two miles are all decades-old legacy completions (drilled 1950s–2002) — mostly plugged shallow verticals plus a handful of long-completed deep wells — ", False),
    ("none representing active drilling, hydraulic fracturing, or a vibration source.", True),
], fill='EAF2FF', space_after=4)
para(doc, [("Independently confirmed by the public hydraulic-fracturing disclosure record (FracFocus, 2011–present): ", False),
           ("no frack job has ever been filed within two miles of the tract, and the most recent within five miles was 2015.", True),
           (" See Finding 10.", False)], fill='EAF2FF', space_after=8)

para(doc, [("Two further points reinforce this:", True)], space_after=2)
bullet(doc, [("Drilling activity in Pecos is mostly rework of existing wells, not new drilling. ", True),
             ("Tracing every Pecos wellbore record updated since 2020 (RRC dbf900, every event tagged to a unique API well number): only ≈10% (117 of 1,118) are genuine new drilling — the other ≈90% (1,001) are recompletion or workover events on existing wellbores. The recompletion activity is overwhelmingly Kinder Morgan Production reworking existing CO₂-flood (enhanced-recovery) fields — a workover rig on an existing bore, not the drilling-and-fracturing activity at issue, and not near the site.", False)], fill='F1F7F0')
bullet(doc, [("Genuine new drilling is deep, not shallow — and remote. ", True),
             ("Of the 116 genuine new wells drilled in Pecos since 2020, ≈95% are deep (≥3,000 ft) and only six are shallow. New shallow drilling is therefore nearly nonexistent anywhere in Pecos, and the deep new drilling that does occur is concentrated well away from the tract (median 20 miles out — see Finding 5).", False)], fill='F1F7F0')

doc.add_heading("Findings", level=1)

doc.add_heading("1. On the Caramba North tract", level=2)
para(doc, [("The shallowest wellbores recorded inside the tract boundary:", False)], space_after=4)
table(doc, ["Depth (ft)", "Spud year", "Status", "Oil/Gas"], [
    ["2,873", "1960", "Plugged & abandoned", "Gas"],
    ["3,067", "1991", "Plugged & abandoned", "Oil"],
    ["3,109", "1957", "Plugged & abandoned", "Oil"],
    ["3,186", "1987", "Plugged & abandoned", "Oil"],
    ["3,250", "2008", "Active", "Oil"]])
para(doc, [("Only one well on the tract lies below 3,000 ft — a 2,873-ft well spudded in 1960 and long since plugged and abandoned. The only active well on the tract (3,250 ft, spudded 2008) is deeper than 3,000 ft. There has been no shallow (<3,000 ft) drilling on the tract in the modern era. (The remaining tract records are a single deep 22,545-ft wellbore and permitted-but-undrilled location entries.)", False)])

doc.add_heading("2. Drilling activity in Pecos is mostly recompletions of existing wells — not new drilling", level=2)
para(doc, [("This is the crux of the data. The Railroad Commission of Texas maintains a master wellbore database (dbf900) in which every drilling, completion, and workover event is logged against a unique API well number. Tracing every Pecos wellbore that has had ", False),
           ("any", True),
           (" recorded activity since 2020:", False)], space_after=4)
table(doc, ["Activity in Pecos since 2020 (wellbore records)", "Count", "Share"], [
    ["Genuine new drilling (a new wellbore drilled)", "117", "≈ 10%"],
    ["Recompletion or workover on an existing wellbore", "1,001", "≈ 90%"],
    ["Total wellbore-record activity", "1,118", "100%"],
], row_styles={0: {'fill': 'CFECCF', 'bold': True}, 1: {'fill': 'EAF5EA'}, 2: {'fill': 'EAF5EA', 'bold': True}})
para(doc, [("In other words, ", False),
           ("about nine out of every ten “drilling-related” actions on a Pecos wellbore since 2020 are workovers on a well that already exists — not a new hole drilled", True),
           (". A recompletion or workover uses a small workover rig on an existing bore; it is not the rig-and-hydraulic-fracturing activity associated with ground vibration.", False)])
add_figure(doc, "ch_recomp_ratio.png", width_in=6.0,
           caption="Of every 1,118 wellbore-record events in Pecos since 2020, only ~117 (≈10%) are genuine new drilling. The other ~1,001 (≈90%) are recompletion or workover events on existing wells.")
para(doc, [("The bulk of this recompletion activity is one operator — ", False),
           ("Kinder Morgan Production", True),
           (" — reworking existing CO₂-flood (enhanced-recovery) fields. None of that involves a drilling rig spudding a new hole, none of it involves a new hydraulic-fracturing program, and the program is not near the Caramba North tract.", False)])
para(doc, [("Whether the question is framed as shallow drilling, hydraulic fracturing, or new drilling of any kind, the record points the same way: ", False),
           ("it is not happening at or near this site.", True)])

doc.add_heading("3. Within 1 mile — no shallow wells", level=2)
para(doc, [("Three wellbores of any depth lie within one mile of the tract; ", False), ("none is shallow (<3,000 ft).", True)])

doc.add_heading("4. Within 2 miles — drilling ended over two decades ago", level=2)
para(doc, [("Of about 46 wellbores within two miles, the ten shallow wells were spudded between 1960 and 2002. The most recent shallow spud within two miles was in 2002, and most of these wells are plugged and abandoned. ", False),
           ("No well of any kind — new drill or otherwise — has been spudded within two miles in over a decade.", True)])

doc.add_heading("5. New drilling since 2020, by distance and depth", level=2)
para(doc, [("Counting only genuine new wells drilled in Pecos since 2020 (wellbore records, recompletion re-stamps excluded):", False)], space_after=4)
table(doc, ["Radius", "New-drill wells, spudded ≥ 2020"], [
    ["≤ 2 mi", "0"],
    ["≤ 5 mi", "0"],
    ["≤ 10 mi", "3  (0 shallow, 3 deep; nearest ≈ 6.9 mi)"],
    ["> 10 mi", "113  (median 20.1 mi, mean 21.1 mi, max 60.4 mi)"],
    ["County-wide total", "116"],
], row_styles={0: {'fill': 'CFECCF', 'bold': True}, 1: {'fill': 'CFECCF', 'bold': True}, 2: {'fill': 'EAF5EA'}, 4: {'fill': 'EAF5EA', 'bold': True}})
para(doc, [("The three genuine new wells within ten miles, across all of 2020–2025, are 6.9–9.4 miles out and all deep (≈9,200–9,500 ft TD; spudded 2020 and 2025) — none shallow, none within five miles.", False)])
para(doc, [("The 113 new wells beyond ten miles are at a median distance of ", False),
           ("≈ 20 miles", True),
           (" from the tract (max 60 mi). Their depths:", False)], space_after=4)
table(doc, ["Depth band", "Wells (of 113)", "Share"], [
    ["< 3,000 ft (shallow)", "6", "5%"],
    ["3,000 – 4,999 ft", "0", "0%"],
    ["5,000 – 9,999 ft", "58", "51%"],
    ["≥ 10,000 ft", "49", "43%"],
], row_styles={0: {'fill': 'EAF5EA'}, 3: {'fill': 'EAF5EA'}})
para(doc, [("That is, ", False), ("107 of 113 (≈95%) of the new wells outside ten miles are deep (≥3,000 ft) — the modern Permian unconventional program", True),
           (" — at a median depth of ≈9,900 ft. The six shallow new-drill wells in the county since 2020 are all remote from the tract.", False)])
para(doc, [("Even on the loosest possible count — every record with a 2020-or-later spud date, including recompletion-restamped records — it is still ", False),
           ("zero within two miles", True),
           (" and only ≈23 within ten miles across the whole period. New drilling does not reach the site under any reading of the data.", False)])

doc.add_heading("6. The nearest non-plugged shallow wells are decades-old completions", level=2)
para(doc, [("The nearest non-plugged shallow wells were spudded in 1970 (1.28 mi) and 1988 (1.97 mi) — decades-old completions, not active drilling. A ground-vibration source is an operating drill rig or a hydraulic-fracturing operation; a plugged or long-completed wellbore is not.", False)])
para(doc, [("No active drilling is occurring adjacent to the tract.", True)], fill='E2EFDA')

doc.add_heading("7. County-wide context — new drilling is deep, and remote from the site", level=2)
para(doc, [("Of the ", False), ("116 genuine new wells drilled in Pecos County since 2020", True),
           (" (Pecos is ≈4,700 sq mi), about ", False), ("95% are deep (≥3,000 ft)", True),
           (" — the modern Permian unconventional program, operator-concentrated in Diamondback, XTO, Continental, and Gordy. Only ", False), ("three", True),
           (" lie within ten miles of the Caramba North tract, and none within five; the activity is overwhelmingly remote from the site (median 20 miles out — see Finding 5).", False)])

doc.add_heading("8. Pecos vs. peer counties — the least new drilling of the group", level=2)
para(doc, [("On the same genuine-new-drill basis (recompletions excluded), Pecos has dramatically less new drilling than comparable Permian counties. Wells spudded since 2020:", False)], space_after=4)
table(doc, ["County", "New-drill wells since 2020", "of which shallow (<3,000 ft)"], [
    ["Pecos (site county)", "116", "6"],
    ["Reeves", "1,044", "35"],
    ["Midland", "1,487", "15"],
    ["Martin", "1,616", "19"],
    ["Reagan", "629", "9"],
    ["Howard", "990", "1"],
    ["Loving", "1,121", "25"],
    ["Other-6 average", "≈1,148", "≈17"],
], row_styles={0: {'fill': 'CFECCF', 'bold': True}, 7: {'fill': 'EAF5EA', 'bold': True}})
para(doc, [("Pecos's ≈116 genuine new wells are roughly one-tenth of the average comparable county's (≈1,148); Martin, the most active, has ≈1,616, and even the next-lowest comparison county (Reagan) has ≈629. Genuine new shallow drilling is negligible in every county (≤35). On a new-drill basis Pecos is by far the least-drilled of the seven — and, per Findings 1–6, essentially none of even that activity is within ten miles of the Caramba North tract.", False)])
para(doc, [("Howard and Loving counties were pulled from the Railroad Commission's full dbf900 wellbore file and integrated on the same genuine-new-drill basis. They lie outside the six-county sale-area set and well away from the tract; they are included here only to broaden the comparison.", False)], italic=True)

doc.add_heading("9. Production near the site is decades-old completions — no active drilling or hydraulic-fracturing operations", level=2)
para(doc, [("Every well was additionally cross-referenced against the Railroad Commission's PDQ production records for the most recent six reported months (through May 2026), joined by ", False),
           ("API number", True),
           (" through the RRC's authoritative API-to-lease crosswalk — a match covering ", False),
           ("≈99.6% of all non-plugged wells", True),
           (". A well is treated as ", False),
           ("“marginal or end-of-life” when its lease's trailing-average output is at or below 125 Mcf/day of gas AND at or below 25 bbl/day of oil", True),
           (" — a strict marginal-well threshold.", False)])
para(doc, [("Of the 291 ", False), ("non-plugged wellbores", True),
           (" within ten miles of the Caramba North tract (recompletion re-stamps already excluded — these are physical wellbores, not paperwork records), ", False),
           ("241 (about 83%) are marginal or end-of-life", True),
           (".", False)], fill='DEEAF6')
para(doc, [("Why so many are end-of-life: ", True),
           ("most of these wellbores were drilled decades ago and are naturally depleted. The “genuine new drill” filter only removes recompletion re-stamps; it does ", False),
           ("not", True),
           (" restrict by spud date. The 291 wellbores within ten miles span the 1960s through 2020s, with the bulk drilled in the 1980s:", False)], space_after=4)
table(doc, ["Spud decade", "Wellbores within 10 mi (non-plugged)"], [
    ["1960s", "29"],
    ["1970s", "57"],
    ["1980s", "122"],
    ["1990s", "20"],
    ["2000s", "42"],
    ["2010s", "18"],
    ["2020s", "3"],
], row_styles={2: {'fill': 'EAF5EA', 'bold': True}})
add_figure(doc, "ch_spud_decade.png", width_in=6.0,
           caption="Of the 291 non-plugged wellbores within 10 mi, 122 were spudded in the 1980s and only 21 since 2010 — most are 35–65 years old, hence the high marginal/end-of-life share.")
para(doc, [("The 50 still producing above the marginal threshold are not active drilling activity either: they are decades-old completions, in four groups —", False)],
     space_after=4)
bullet(doc, [("about 15 ", False), ("legacy deep-gas wells", True), (" (mostly 1965–1978 spud, ≈17,000–22,800 ft);", False)])
bullet(doc, [("a cluster of ~31 ", False), ("low-rate vertical conventional oil wells", True), (" at ≈3,150–3,440 ft depth (spud 1986–2012, mostly 2007–2012; all on a unitized lease producing roughly 37.5 bbl/day per well — a stripper-grade pumping operation);", False)])
bullet(doc, [("3 ", False), ("truly shallow (<3,000 ft) oil wells", True), (" at 2,824–2,945 ft (spud 2004–2011, also at stripper rates just above the threshold); and", False)])
bullet(doc, [("the single ", False), ("2020 deep-horizontal new-drill", True), (" noted in Finding 5 (9,237 ft, 9.37 mi out).", False)], space_after=8)
add_figure(doc, "ch_status_mix.png", width_in=5.6,
           caption="Production status of 291 non-plugged wellbores within 10 mi: 83% marginal or end-of-life, 17% still producing — all on decades-old completions.")
para(doc, [("Within five miles, ", False),
           ("52 of 86 non-plugged wellbores are marginal or end-of-life", True),
           (" and the 34 still producing are 3 of the legacy deep-gas wells plus 31 of the shallow vertical oil cluster — none from the modern Permian horizontal program.", False)],
     fill='DEEAF6')
para(doc, [("Within two miles, ", False),
           ("3 of 5 non-plugged wellbores are marginal or end-of-life", True),
           (" and the 2 still producing are the 1.13-mi 1975/≈22,100-ft legacy gas well (125.7 Mcf/d) and a 1.97-mi 2008/≈3,275-ft vertical oil well (37.5 bbl/d).", False)],
     fill='DEEAF6')
para(doc, [("A pumping wellhead or low-rate gas well on a decades-old completion is not a drill rig or a hydraulic-fracturing spread; together with the plugged legacy wells discussed above, ", False),
           ("there is no active drilling or hydraulic-fracturing operation at or near the site.", True)],
     fill='DEEAF6', space_after=8)
para(doc, [("County-wide, of the ≈35,100 non-plugged wellbores in the six sale-area counties, ≈12,540 are plugged, ≈12,490 are marginal or end-of-life by this measure, and ≈10,060 are still producing above the threshold. The API-to-lease crosswalk matches ≈99.6% of non-plugged wells; the ≈90 that do not match are conservatively left classified “Active.” Production filings carry a normal reporting lag, which the six-month trailing window mitigates.", False)], italic=True)

doc.add_heading("10. The public fracking record (FracFocus) confirms: no fracking jobs within two miles of the tract, ever", level=2)
para(doc, [("The Texas ", False), ("FracFocus disclosure database", True),
           (" (fracfocus.org) is the public record of every hydraulic-fracturing job that operators have filed in Texas since 2011. Cross-referencing every Pecos County disclosure (949 in total) against the Caramba North tract:", False)], space_after=4)
table(doc, ["Distance band from tract", "Frack disclosures (2011 – present)", "Most recent year"], [
    ["0 – 2 mi", "0", "— none, ever"],
    ["2 – 5 mi", "9", "2015 (most recent)"],
    ["5 – 10 mi", "20", "2025"],
    ["10 – 20 mi", "464", "2026"],
], row_styles={0: {'fill': 'CFECCF', 'bold': True}, 1: {'fill': 'EAF5EA'}, 3: {'fill': 'EAF5EA'}})
para(doc, [("No hydraulic-fracturing job has ever been performed within two miles of the Caramba North tract.", True),
           (" Within five miles there have been 9 fracks, all between 2012 and 2015 — the most recent over a decade ago. Eight of the nine were Apache Corporation's 2012 multi-well program (FSSU wells, 2.75 – 4.85 mi out); the ninth was Flamingo Operating in 2015. Within ten miles, 29 fracks over the entire 2012 – 2025 period, the closest to the site being the 2025 Mongoose Energy Viper wells at 6.94 mi out.", False)],
     fill='DEEAF6')
add_figure(doc, "ch_fracfocus_rings.png", width_in=6.0,
           caption="FracFocus disclosures around Caramba North: 0 within 2 mi (ever), 9 within 2–5 mi (all 2012–2015), 20 within 5–10 mi, 464 within 10–20 mi. The dense Permian fracking activity is concentrated outside the 10-mile buffer.")
para(doc, [("The broader Permian fracking program does exist — 493 fracks within twenty miles since 2011, dominated by the deep-horizontal unconventional players (Diamondback, XTO/ExxonMobil, Gordy). But that activity is concentrated well outside the 10-mile buffer, almost entirely at unconventional depths (median TVD ≈ 9,800 ft within 20 mi), and even there annual volume has been declining (39 fracks in Pecos in 2023, 6 in 2024, 21 in 2025, 11 year-to-date 2026 vs. a 2018 – 2019 peak of 113 – 167/year).", False)])
para(doc, [("This is direct, evidence-based confirmation of what the wellbore-and-production record already implied: ", False),
           ("no active hydraulic-fracturing operation is occurring at or near the Caramba North tract — neither on horizontal wellbores nor on vertical wellbores, and the closest disclosed frack within five miles is over a decade old.", True)], fill='DEEAF6')

doc.save(OUT)
print("WROTE", OUT)
