#!/usr/bin/env python3
"""
Build the condensed "Executive Brief" for Caramba North (redesign brief §7,
Doc 1) as a Word document using python-docx.

Portrait, US Letter, 5-6 page hard ceiling. Memo style: headline + insight
subheading + short paragraph + one supporting number/exhibit per topic.
No tables — prose only, numbers inline / as pull-out stat lines.

Facts and copy are sourced from docs/redesign_content_brief.md — do not
invent figures. Distances are the corrected edge-to-edge figures (GW Ranch
~15.5 mi, Longfellow ~19.3 mi), never the old centroid figures.

Usage:
    python3 scripts/build_brief_executive.py
"""
import os
import subprocess
import json

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EXHIBIT_DIR = os.path.join(REPO, "outputs", "reports", "om_exhibits")
OUT = os.path.join(REPO, "outputs", "reports", "Caramba-North-Brief-Executive.docx")

# Pull the corrected distances / ring analysis from the shared insight pack
# to cross-check the hardcoded copy below (kept in sync with the brief).
TMP_JSON = "/tmp/om_brief_exec.json"
subprocess.run(
    ["python3", os.path.join(REPO, "scripts", "build_insight_pack.py"), "--json", TMP_JSON],
    cwd=REPO, check=True, capture_output=True,
)
with open(TMP_JSON) as f:
    PACK = json.load(f)
assert PACK["distances_edge_to_edge"]["gw_ranch_mi"] == 15.5
assert PACK["distances_edge_to_edge"]["longfellow_mi"] == 19.3

# ------------------------------------------------------------------ palette
NAVY = RGBColor(0x0F, 0x1B, 0x2D)
SLATE = RGBColor(0x33, 0x41, 0x52)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RED = RGBColor(0xB9, 0x1C, 0x1C)
LINE = RGBColor(0xCB, 0xD5, 0xE1)
FAINT = RGBColor(0x94, 0xA3, 0xB8)

HEADLINE_FONT = "Cambria"
BODY_FONT = "Calibri"

# ------------------------------------------------------------------ doc setup
doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.different_first_page_header_footer = True
section.header_distance = Inches(0.4)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = SLATE
normal.paragraph_format.space_after = Pt(0)
normal_rpr = normal.element.get_or_add_rPr()
rfonts = normal_rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    normal_rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), BODY_FONT)


PPR_CHILD_ORDER = [
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
    "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
]
PPR_ORDER_QNS = [qn(t) for t in PPR_CHILD_ORDER]


def insert_pPr_child(pPr, child):
    """Insert child into pPr respecting CT_PPrBase's fixed element sequence
    (python-docx's paragraph_format setters append blindly, which produces
    schema-invalid ordering once a pBdr/ind is added after spacing exists)."""
    idx = PPR_ORDER_QNS.index(child.tag)
    for existing in list(pPr):
        if existing.tag in PPR_ORDER_QNS and PPR_ORDER_QNS.index(existing.tag) > idx:
            existing.addprevious(child)
            return
    pPr.append(child)


def set_field(paragraph, instr):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    run._r.addnext(fld)
    r_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.text = "1"
    r_el.append(t_el)
    fld.append(r_el)
    run._r.getparent().remove(run._r)


# First page (cover) header/footer left blank
doc.sections[0].first_page_header.paragraphs[0].text = ""
doc.sections[0].first_page_footer.paragraphs[0].text = ""

# Running header (pages 2+)
hdr_p = section.header.paragraphs[0]
hdr_p.text = ""
r = hdr_p.add_run("CARAMBA NORTH  ·  EXECUTIVE BRIEF")
r.font.size = Pt(8)
r.font.color.rgb = MUTED
r.font.name = BODY_FONT
r.font.bold = True
hdr_p.paragraph_format.space_after = Pt(2)
# bottom border under header text
pPr = hdr_p._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "4")
bottom.set(qn("w:space"), "4")
bottom.set(qn("w:color"), "CBD5E1")
pbdr.append(bottom)
insert_pPr_child(pPr, pbdr)

# Running footer (pages 2+): CONFIDENTIAL left, page number right via tab
ftr_p = section.footer.paragraphs[0]
ftr_p.text = ""
tab_stops = ftr_p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Inches(6.8))
r1 = ftr_p.add_run("CONFIDENTIAL — PREPARED UNDER NDA")
r1.font.size = Pt(7.5)
r1.font.color.rgb = FAINT
r1.font.name = BODY_FONT
r2 = ftr_p.add_run("\tPage ")
r2.font.size = Pt(7.5)
r2.font.color.rgb = FAINT
r2.font.name = BODY_FONT
set_field(ftr_p, "PAGE")
ftr_p.runs[-1].font.size = Pt(7.5)
ftr_p.runs[-1].font.color.rgb = FAINT
ftr_p.runs[-1].font.name = BODY_FONT


# ------------------------------------------------------------------ helpers
def para(spacing_before=0, spacing_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    if align:
        p.alignment = align
    return p


def run(p, text, size=10.5, bold=False, italic=False, color=SLATE, font=BODY_FONT, caps=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = font
    if caps:
        r.font.all_caps = True
    return r


def hr(color="CBD5E1", size_pt=0.75, space_after=10, space_before=2):
    """Thin horizontal rule used to separate stacked topic blocks on a page."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_pt * 8)))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    insert_pPr_child(pPr, pbdr)
    return p


def headline(text, size=19, color=NAVY, space_before=0, space_after=3):
    p = para(spacing_before=space_before, spacing_after=space_after)
    run(p, text, size=size, bold=True, color=color, font=HEADLINE_FONT)
    return p


def subheading(text, color=SLATE, size=11.5, space_after=7):
    p = para(spacing_after=space_after)
    run(p, text, size=size, italic=True, bold=False, color=color, font=HEADLINE_FONT)
    return p


def body(text, size=10.3, space_after=7, color=SLATE):
    p = para(spacing_after=space_after)
    run(p, text, size=size, color=color)
    return p


def stat_line(label, color=NAVY, space_after=13):
    """A short bold pull-out number/fact — the 'one supporting number' per topic."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "B91C1C" if color == RED else "0F1B2D")
    pbdr.append(left)
    insert_pPr_child(pPr, pbdr)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "144")
    insert_pPr_child(pPr, ind)
    run(p, label, size=11.5, bold=True, color=color, font=BODY_FONT)
    return p


def add_image(path, width_in, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run()
    r.add_picture(path, width=Inches(width_in))
    return p


def caption(text, size=8, space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = para(spacing_after=space_after, align=align)
    run(p, text, size=size, italic=True, color=MUTED)
    return p


def small_print(text, size=7.6, color=FAINT, space_after=6):
    p = para(spacing_after=space_after)
    run(p, text, size=size, color=color, font=BODY_FONT)
    return p


def topic_block(title, subhead, paragraph_text, stat_text, stat_color=NAVY,
                 title_color=NAVY, subhead_color=SLATE, first=False):
    if not first:
        hr()
    headline(title)
    subheading(subhead, color=subhead_color)
    body(paragraph_text)
    stat_line(stat_text, color=stat_color)


# =============================================================== COVER PAGE
p = para(spacing_before=6, spacing_after=2)
run(p, "CARAMBA NORTH", size=34, bold=True, color=NAVY, font=HEADLINE_FONT)

p = para(spacing_after=4)
run(p, "EXECUTIVE BRIEF", size=14, bold=True, color=RED, font=BODY_FONT, caps=True)

p = para(spacing_after=16)
run(p, "Confidential — prepared for a limited number of prospective counterparties under NDA",
    size=8.5, italic=True, color=MUTED)

positioning = (
    "Caramba North is a 1,300-acre parcel sitting inside an already-forming power and "
    "data-center corridor. Two hyperscale-scale projects — 7.65 GW and 2 GW — sit "
    "on the same north-south line through the property, at 15.5 and 19.3 miles, backed by "
    "a transmission, water, and gas position that is already permitted, not proposed. "
    "The region’s own numbers do the selling: 32.7 GW of operating and queued power "
    "capacity within 60 miles, in a state where the interconnection backlog has grown large "
    "enough — 474 GW, about 90% data centers — to trigger a gubernatorial audit and a "
    "queue-processing pause. Caramba North is positioned to benefit from the same buildout "
    "without carrying the exposure of the marginal, unproven project in that queue: water "
    "and gas are already under contract-ready terms, not just an application."
)
p = para(spacing_after=14)
run(p, positioning, size=11.2, color=SLATE)

add_image(os.path.join(EXHIBIT_DIR, "exhibit_power_gravity_light.png"), 4.35, space_after=4)
caption(
    "Operating + ERCOT-queued power capacity by radius from Caramba North, with GW Ranch "
    "and Longfellow plotted at true bearing and distance.",
    space_after=10,
)
stat_line("32.7 GW operating + queued within 60 miles of Caramba North", color=NAVY, space_after=0)

doc.add_page_break()

# =============================================================== PAGE 2: Property + Transmission
topic_block(
    "The Property",
    "As-of-right industrial land inside the fastest-growing load pocket in ERCOT — not a rezoning story.",
    "The tract totals 1,300 acres of maximum contiguous acreage on the north side of "
    "Interstate 10 in Pecos County, roughly five miles from Fort Stockton’s services and "
    "regional airport. The site carries no zoning ordinance, so industrial and energy use is "
    "available as of right — there is no rezoning process standing between acquisition and "
    "construction. It sits in ERCOT’s Far West weather zone, the system’s highest-growth "
    "large-load pocket.",
    "1,300 contiguous acres, as-of-right industrial/energy use",
    first=True,
)

topic_block(
    "The Transmission Position",
    "Fifteen miles from the delivery point of all three approved 765 kV Permian import lines — "
    "the transmission decision is already made, upstream of this site.",
    "The property sits 15 miles from AEP/CPS Energy’s Solstice Substation, the western "
    "terminus of all three PUCT-approved 765 kV Permian import paths (approved April 24, "
    "2025, PUCT Docket No. 55718). Six local substations sit within 10 miles of the tract. "
    "ERCOT is separately tracking 141 substation and 133 line upgrades system-wide under its "
    "Transmission Planning Improvement Tool — a pipeline of planned upgrades, not yet built, "
    "offered here as corridor context rather than committed capacity.",
    "15 mi to Solstice Substation — terminus of all three approved 765 kV Permian lines",
)

doc.add_page_break()

# =============================================================== PAGE 3: Power Cluster + Water/Gas
topic_block(
    "The Regional Power Cluster",
    "12 GW already queued in this county alone — before counting the two hyperscale "
    "campuses profiled on the next page.",
    "Pecos County has 3,226 MW of generation operating today — 2,178 MW solar, 542 MW "
    "wind, 505 MW battery storage, and 1 MW gas — plus 12,039 MW queued across 39 ERCOT "
    "interconnection projects. The six adjacent counties (Reeves, Crane, Ward, Upton, Ector, "
    "Crockett) add 7,022 MW operating and 24,585 MW queued. Within 20 miles of the tract "
    "specifically, 13 queued projects total 3,973 MW, and the nearest operating storage "
    "asset, St. Gall Energy Storage I, sits 1.9 miles away.",
    "12,039 MW queued in Pecos County alone (39 ERCOT projects)",
    first=True,
)

topic_block(
    "Water & Gas",
    "Two-thirds of the district’s industrial water rights are already permitted to this "
    "position — the water conversation is closed, not open.",
    "47,418 acre-feet per year (42.3 MGD) is permitted on adjacent affiliated lands, drawn "
    "from the Edwards-Trinity (Plateau) aquifer, whose recharge held through the 1950s "
    "drought of record — roughly two-thirds of all Middle Pecos Groundwater Conservation "
    "District industrial water rights. On gas, the site sits 20 miles from the Waha hub, "
    "where an indicative supply quote is already in hand: 200,000 MMBtu/day on a 15-year "
    "term at Waha-index pricing, with a $15–25 million CIAC and a 9–15 month lead time. "
    "Waha has priced at a structural discount to Henry Hub, including negative prints in "
    "2024–2025 as new pipelines rebalance Permian egress.",
    "47,418 AF/yr permitted water · 200,000 MMBtu/day gas quote in hand",
)

doc.add_page_break()

# =============================================================== PAGE 4: Regional Pipeline
headline("The Regional Data-Center Pipeline")
subheading(
    "GW Ranch sits almost due north and Longfellow almost due south of Caramba North — "
    "the property sits on the line between them, not off to one side."
)
body(
    "Two hyperscale-scale anchors bracket the property on the same north-south corridor. "
    "Across both, 79.3% of combined announced capacity (GW Ranch) is under construction and "
    "20.7% (Longfellow) is in planned/phase-1 status — the regional pipeline is majority-"
    "built, not majority-speculative.",
    space_after=7,
)

p = para(spacing_after=2)
run(p, "GW Ranch (Amazon)", size=13, bold=True, color=RED, font=HEADLINE_FONT)
subheading(
    "The largest air permit issued in the US this year sits fifteen miles up the same "
    "highway corridor — under construction, not announced.",
    color=SLATE, size=10.5, space_after=4,
)
body(
    "GW Ranch is an 8,000-acre site in Pecos County, approximately 15.5 miles from Caramba "
    "North (edge-to-edge), disclosed as Amazon-owned in August 2026 (previously developed by "
    "Pacifico Energy Group, which remains the power-plant developer/operator). It holds a "
    "7.65 GW TCEQ air permit for 35 gas turbines — issued January/February 2026 and the "
    "largest in the US — plus 1.8 GW of battery storage and up to 750 MW of solar, with "
    "three data-center buildings targeted for completion in December 2026 and an estimated "
    "$12 billion in total project investment. The 7.65 GW figure is a TCEQ generation air "
    "permit, not an ERCOT interconnection queue position; the project is off-grid initially.",
    size=9.6, space_after=5,
)
add_image(os.path.join(EXHIBIT_DIR, "exhibit_amz_gwranch.jpg"), 4.4, space_after=2)
stat_line("7.65 GW air permit — largest in the US · ≈15.5 mi from Caramba North", color=RED, space_after=9)

p = para(spacing_after=2)
run(p, "Longfellow", size=13, bold=True, color=RED, font=HEADLINE_FONT)
subheading(
    "A second phased gas-generation campus twenty miles south — the corridor’s demand "
    "for on-site power isn’t one project deep.",
    color=SLATE, size=10.5, space_after=4,
)
body(
    "Longfellow is a 568-acre site in Pecos County, approximately 19.3 miles from Caramba "
    "North (edge-to-edge); the site’s own public materials describe the location as more "
    "than 25 miles outside Fort Stockton. On-site natural-gas generation is planned using "
    "aero-derivative turbines with SCR and carbon-capture capability, cooled on a closed "
    "loop using permitted non-potable groundwater. Originally announced in October 2025 as a "
    "2 GW, eight-phase campus, phase-1 site work is underway, with the on-site generation "
    "build planned in phases. No confirmed ERCOT queue position or TCEQ air-permit record has "
    "been found for this site as of August 2026.",
    size=9.6, space_after=5,
)
add_image(os.path.join(EXHIBIT_DIR, "exhibit_longfellow.jpg"), 4.4, space_after=2)
stat_line("2 GW planned, 8 phases · ≈19.3 mi from Caramba North", color=RED, space_after=0)

doc.add_page_break()

# =============================================================== PAGE 5: Subsurface + Notices
topic_block(
    "Subsurface & Drilling Activity",
    "Pecos County has the lowest new-drilling count of seven comparable Permian counties "
    "since 2020 — a 90%-below-peer-average level of activity, not merely quiet.",
    "Pecos County has recorded 115 new-drill wellbore events since 2020, out of 1,140 total "
    "RRC events — 90% of recorded activity is workovers or reworks, not new drilling. Zero "
    "new-drill wells sit within 2 miles of the tract since 2020, zero within 5 miles, and "
    "only one within 10 miles (9.37 miles away); beyond 10 miles, the median distance to the "
    "remaining 114 wells is 19.9 miles. Against six comparable Permian counties — Reagan "
    "(668), Howard (990), Reeves (1,053), Loving (1,121), Midland (1,569), and Martin (1,685) "
    "— Pecos’s 115 new-drill count is the lowest of all seven, roughly 90% below the "
    "1,181 peer average. Within 10 miles, 83% of non-plugged wellbores are marginal or "
    "end-of-life production, an even higher share than the 60–62% seen closer in.",
    "115 new-drill wells since 2020 — lowest of 7 peer counties (vs. 1,181 average)",
    first=True,
)

hr(space_after=8)
p = para(spacing_after=4)
run(p, "NOTICES", size=8.5, bold=True, color=MUTED, caps=True)
small_print(
    "This is a confidential offering memorandum prepared for a limited number of "
    "prospective counterparties under NDA. It is not an offer to sell or a solicitation of "
    "securities. Information is preliminary and indicative, from sources believed reliable. "
    "Public data is drawn from the sources cited in the underlying diligence platform "
    "(ERCOT GIS Report/TPIT, PUCT, EIA-860, TCEQ, RRC, FracFocus, Middle Pecos GCD, HIFLD, "
    "USGS, BTS, Census TIGER); third-party transaction news is sourced to public reporting "
    "cited in the companion source register.",
    space_after=8,
)

p = para(spacing_after=4)
run(p, "DISTANCE METHODOLOGY", size=8.5, bold=True, color=MUTED, caps=True)
small_print(
    "Distances to GW Ranch and Longfellow are measured edge-to-edge — from the nearest "
    "point on the Caramba North tract boundary to each site’s disclosed location — rather "
    "than centroid-to-centroid, which is consistently longer because the Caramba North tract "
    "itself has spatial extent (GW Ranch: 15.5 mi vs. 17.3 mi centroid; Longfellow: 19.3 mi "
    "vs. 19.7 mi centroid). Longfellow’s own public site describes its location as more than "
    "25 miles outside Fort Stockton, consistent with the longer figure; this distance should "
    "not be represented as shorter.",
    space_after=0,
)

# python-docx's default settings.xml template emits <w:zoom w:val="bestFit"/>
# without the required w:percent attribute (schema-invalid on its own,
# independent of anything else in this script) — patch it directly.
zoom_el = doc.settings.element.find(qn("w:zoom"))
if zoom_el is not None and zoom_el.get(qn("w:percent")) is None:
    zoom_el.set(qn("w:percent"), "100")

doc.save(OUT)
print("wrote", OUT)
