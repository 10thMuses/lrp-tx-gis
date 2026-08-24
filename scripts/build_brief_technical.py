#!/usr/bin/env python3
"""Build the Caramba North "Technical Snapshot" condensed document
(brief §7, Doc 2) — a dense, table/stat-grid driven counterpart to the
prose-based "Executive Brief". python-docx per the docx skill's guidance.

Facts and copy are sourced from docs/redesign_content_brief.md. Ring
analysis / distances / project-maturity figures come from
scripts/build_insight_pack.py (uses a unique tmp json path).

Usage:
    python3 scripts/build_brief_technical.py
"""
from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Emu

REPO = Path(__file__).resolve().parent.parent
EXHIBIT_DIR = REPO / "outputs" / "reports" / "om_exhibits"
OUT = REPO / "outputs" / "reports" / "Caramba-North-Brief-Technical.docx"
TMP_JSON = "/tmp/om_brief_tech.json"

# ---------------------------------------------------------------------------
# Rule-3 fix-up for the two shared, do-not-regenerate exhibit rasters.
# Both exhibit_amz_gwranch.jpg and exhibit_longfellow.jpg carry a baked-in
# "Longfellow — Project Horizon · Poolside / CoreWeave · 2 GW U/C" map label
# (CoreWeave is the named tenant) that violates brief Rule 3 for this
# document. This does not touch or regenerate the shared source rasters or
# the map logic that produced them (brief §5/§7a); it derives build-specific
# crops into our own subfolder, same technique already used by the
# Minimalist deck's build_minimalist_exhibit_crops.py for the identical
# problem, so concurrent builds reading the shared files are unaffected.
_FIXED_EXHIBIT_DIR = EXHIBIT_DIR / "_technical"
_FIXED_EXHIBIT_DIR.mkdir(exist_ok=True)


def _fixup_exhibits():
    from PIL import Image, ImageDraw

    # GW Ranch: crop off the bottom strip carrying the incidental Longfellow/
    # CoreWeave label (unrelated to the GW Ranch story on this page).
    src = Image.open(EXHIBIT_DIR / "exhibit_amz_gwranch.jpg")
    w, h = src.size  # 1410 x 540
    cropped = src.crop((0, 0, w, int(h * 450 / 540)))
    gwranch_out = _FIXED_EXHIBIT_DIR / "exhibit_amz_gwranch_notenant.jpg"
    cropped.save(gwranch_out, quality=92)

    # Longfellow: the tenant clause sits inside the highlight ring/label box
    # itself (not croppable without losing the ring, badge, and distance
    # line), so paint over just "Poolside / CoreWeave · 2 GW U/C", in white,
    # flush inside the label box border — leaving "Longfellow — Project
    # Horizon" and every other element untouched.
    src2 = Image.open(EXHIBIT_DIR / "exhibit_longfellow.jpg").convert("RGB")
    draw = ImageDraw.Draw(src2)
    draw.rectangle((393, 101, 833, 139), fill=(255, 255, 255))
    longfellow_out = _FIXED_EXHIBIT_DIR / "exhibit_longfellow_notenant.jpg"
    src2.save(longfellow_out, quality=92)

    return gwranch_out, longfellow_out


GWRANCH_IMG, LONGFELLOW_IMG = _fixup_exhibits()

# ---------------------------------------------------------------- data model
subprocess.run(
    [sys.executable, str(REPO / "scripts" / "build_insight_pack.py"), "--json", TMP_JSON],
    cwd=REPO, check=True,
)
PACK = json.loads(Path(TMP_JSON).read_text())
RING = PACK["ring_analysis"]
MATURITY = PACK["project_maturity"]
DIST = PACK["distances_edge_to_edge"]

STAMP = "August 2026"

# ---------------------------------------------------------------- palette
NAVY = "0F1B2D"
STEEL = "1B4965"          # primary accent — this doc's own accent (not navy, not red)
STEEL_DK = "12324A"
TINT = "EAF1F5"           # steel tint for header bars/callouts
ROW_ALT = "F5F7F8"
GRAPHITE = "26313D"
MUTED = "5B6672"
LINE = "C9D2D8"
RED = "B91C1C"             # reserved for the two feature call-outs only
RED_TINT = "FBEAEA"
WHITE = "FFFFFF"

FONT_SANS = "Arial"
FONT_MONO = "Consolas"

CONTENT_W = 7.4  # inches, at 0.55in margins on 8.5in page


# ============================================================ low-level oxml
# CT_TcPr child sequence per ECMA-376 — required to keep every tcPr valid
# regardless of the order our helpers (or python-docx's own vAlign setter)
# touch a given cell in.
_TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd",
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
    "headers", "cellIns", "cellDel", "cellMerge", "tcPrChange",
]


def _tcpr_place(tcPr, el):
    """Insert el into tcPr at its schema-correct position, replacing any
    existing element of the same tag."""
    tag = el.tag.split("}")[-1]
    idx = _TCPR_ORDER.index(tag)
    existing = tcPr.find(qn(f"w:{tag}"))
    if existing is not None:
        tcPr.remove(existing)
    insert_before = None
    for child in tcPr:
        ctag = child.tag.split("}")[-1]
        if ctag in _TCPR_ORDER and _TCPR_ORDER.index(ctag) > idx:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(el)
    else:
        tcPr.append(el)


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _tcpr_place(tcPr, shd)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    # CT_TcBorders sequence: top, start, left, bottom, end, right, ...
    for tag, spec in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        if spec is None:
            el.set(qn("w:val"), "nil")
        else:
            sz, color = spec
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    _tcpr_place(tcPr, borders)


def _set_cell_margins(cell, top=40, bottom=40, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    # CT_TcMar sequence: top, start, left, bottom, end, right
    for tag, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    _tcpr_place(tcPr, mar)


_TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption",
    "tblDescription",
]


def _tblpr_place(tblPr, el):
    tag = el.tag.split("}")[-1]
    idx = _TBLPR_ORDER.index(tag)
    existing = tblPr.find(qn(f"w:{tag}"))
    if existing is not None:
        tblPr.remove(existing)
    insert_before = None
    for child in tblPr:
        ctag = child.tag.split("}")[-1]
        if ctag in _TBLPR_ORDER and _TBLPR_ORDER.index(ctag) > idx:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(el)
    else:
        tblPr.append(el)


def _fixed_layout(table):
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    _tblpr_place(tblPr, layout)


def _set_grid(table, widths_in):
    tbl = table._tbl
    old = tbl.find(qn("w:tblGrid"))
    if old is not None:
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        grid.append(gc)
    tbl.tblPr.addnext(grid)  # tblGrid must immediately follow tblPr
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])


def _row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    trPr.append(el)


def _row_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _no_borders_table(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, None, None, None, None)
            _set_cell_margins(cell, 0, 0, 0, 0)


def _cell_para(cell):
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _run(para, text, size=8, bold=False, italic=False, color=GRAPHITE, font=FONT_SANS,
          caps=False, spacing=None):
    r = para.add_run(text.upper() if caps else text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = font
    r.font.color.rgb = RGBColor.from_string(color)
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rPr.append(sp)
    return r


def set_col_widths(table, widths_in):
    _set_grid(table, widths_in)
    _fixed_layout(table)


# ============================================================ building blocks
def section_tag_and_title(container, tag, title, space_before=10, space_after=2, width=CONTENT_W):
    """Small colored eyebrow bar: 'SEC 2.1' then the section title, with a
    rule underneath (drawn via a bordered table cell, not a paragraph border,
    to avoid CT_PPr child-ordering constraints)."""
    _spacer(container, space_before if space_before else 1)
    tbl = container.add_table(rows=1, cols=1)
    set_col_widths(tbl, [width])
    cell = tbl.cell(0, 0)
    _set_cell_margins(cell, top=0, bottom=30, left=0, right=0)
    _set_cell_border(cell, bottom=(6, STEEL), top=None, left=None, right=None)
    p = _cell_para(cell)
    _run(p, f"{tag}  ", size=8, bold=True, color=STEEL, font=FONT_MONO, caps=True, spacing=10)
    _run(p, title, size=12.5, bold=True, color=NAVY, font=FONT_SANS)
    _spacer(container, space_after if space_after else 1)
    return p


def insight_callout(container, text, space_after=6, width=CONTENT_W):
    """The brief's 'subheading candidate' line, as a tinted callout bar sitting
    directly under a section title — the conclusion-not-a-label statement,
    rendered as data-grid furniture rather than prose."""
    tbl = container.add_table(rows=1, cols=1)
    set_col_widths(tbl, [width])
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, TINT)
    _set_cell_border(cell, bottom=None, top=None, right=None, left=(18, STEEL))
    _set_cell_margins(cell, top=60, bottom=60, left=140, right=140)
    p = _cell_para(cell)
    _run(p, text, size=9, italic=True, bold=True, color=STEEL_DK)
    _row_cant_split(tbl.rows[0])
    _spacer(container, space_after)
    return tbl


def _spacer(container, pts=6):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pts)
    r = p.add_run("")
    r.font.size = Pt(1)
    return p


def fact_table(container, rows, col_widths, aligns=None, bold_first_col=False,
               feature=False, header=True):
    """rows[0] is the header row when header=True. Dense two-tone data table."""
    accent = RED if feature else STEEL
    accent_tint = RED_TINT if feature else TINT
    n_cols = len(col_widths)
    aligns = aligns or ["left"] * n_cols
    table = container.add_table(rows=len(rows), cols=n_cols)
    set_col_widths(table, col_widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(rows):
        is_header = header and r_idx == 0
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)
            p = _cell_para(cell)
            p.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
            }[aligns[c_idx]]
            if is_header:
                _set_cell_bg(cell, accent)
                _set_cell_border(cell, bottom=(4, accent), top=None, left=None, right=None)
                _run(p, str(val), size=7.2, bold=True, color=WHITE, caps=True, spacing=8)
            else:
                zebra = ROW_ALT if (r_idx % 2 == (0 if not header else 1)) else WHITE
                _set_cell_bg(cell, zebra)
                bottom_color = LINE
                _set_cell_border(cell, bottom=(3, bottom_color), top=None, left=None, right=None)
                is_bold = bold_first_col and c_idx == 0
                _run(p, str(val), size=8.3, bold=is_bold,
                     color=(STEEL_DK if is_bold else GRAPHITE),
                     font=(FONT_MONO if aligns[c_idx] == "right" else FONT_SANS))
        _row_cant_split(table.rows[r_idx])
    if header:
        _row_repeat_header(table.rows[0])
    return table


def two_col(container, left_w=None, right_w=None, gutter=0.25):
    left_w = left_w if left_w is not None else (CONTENT_W - gutter) / 2
    right_w = right_w if right_w is not None else (CONTENT_W - gutter) / 2
    table = container.add_table(rows=1, cols=3)
    set_col_widths(table, [left_w, gutter, right_w])
    _no_borders_table(table)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # clear the default empty paragraph's spacing
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    return table.rows[0].cells[0], table.rows[0].cells[2]


def stat_tile_row(container, tiles):
    """tiles: list of (big_number, label) tuples — the headline strip."""
    n = len(tiles)
    gutter = 0.12
    w = (CONTENT_W - gutter * (n - 1)) / n
    table = container.add_table(rows=1, cols=n)
    set_col_widths(table, [w] * n)
    for i, (num, label) in enumerate(tiles):
        cell = table.cell(0, i)
        _set_cell_bg(cell, NAVY)
        _set_cell_margins(cell, top=110, bottom=110, left=90, right=90)
        _set_cell_border(cell, bottom=(10, STEEL), top=None, left=None, right=None)
        p = _cell_para(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, num, size=15, bold=True, color=WHITE, font=FONT_MONO)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(0)
        _run(p2, label, size=6.6, bold=True, color="C7D7E0", caps=True, spacing=6)
    return table


def small_print(container, text, space_before=4, space_after=2, color=MUTED, size=6.6):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    _run(p, text, size=size, color=color, italic=True)
    return p


def add_picture_fit(container, path, max_w_in, max_h_in=None):
    from PIL import Image
    with Image.open(path) as im:
        iw, ih = im.size
    ar = iw / ih
    w = max_w_in
    h = w / ar
    if max_h_in is not None and h > max_h_in:
        h = max_h_in
        w = h * ar
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(w), height=Inches(h))
    return p


def caption(container, text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=6.8, italic=True, color=MUTED)
    return p


def page_footer(section, label):
    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(CONTENT_W), alignment=2)  # right tab
    _run(p, "CARAMBA NORTH — TECHNICAL SNAPSHOT  |  CONFIDENTIAL", size=6.6, color=MUTED, caps=False)
    r = p.add_run("\t")
    field_p = r
    # page number field
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld_sep = OxmlElement('w:fldChar'); fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run2 = p.add_run()
    run2.font.size = Pt(6.6)
    run2.font.color.rgb = RGBColor.from_string(MUTED)
    run2._r.append(fld_begin); run2._r.append(instr); run2._r.append(fld_sep); run2._r.append(fld_end)


# ============================================================ document setup
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.5)
sec.bottom_margin = Inches(0.55)
sec.left_margin = Inches(0.55)
sec.right_margin = Inches(0.55)
sec.header_distance = Inches(0.25)
sec.footer_distance = Inches(0.3)

style = doc.styles["Normal"]
style.font.name = FONT_SANS
style.font.size = Pt(8.3)
style.font.color.rgb = RGBColor.from_string(GRAPHITE)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

page_footer(sec, "Caramba North")

# ---------------------------------------------------------------- PAGE 1 ---
# Header block
htbl = doc.add_table(rows=1, cols=2)
set_col_widths(htbl, [4.6, CONTENT_W - 4.6])
_no_borders_table(htbl)
lc, rc = htbl.rows[0].cells
p = _cell_para(lc)
_run(p, "CARAMBA NORTH", size=22, bold=True, color=NAVY, font=FONT_SANS)
p2 = lc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
_run(p2, "TECHNICAL SNAPSHOT", size=11, bold=True, color=STEEL, caps=True, spacing=22)
p3 = rc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_run(p3, "CONFIDENTIAL OFFERING MEMORANDUM", size=7, bold=True, color=MUTED, caps=True, spacing=8)
p4 = rc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p4.paragraph_format.space_before = Pt(1)
_run(p4, f"Pecos County, Texas  |  {STAMP}", size=7.4, color=MUTED)
p5 = rc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p5.paragraph_format.space_before = Pt(1)
_run(p5, "Data-room reference — figures per §2.7 diligence platform", size=6.6, italic=True, color=MUTED)

# rule under header (bordered table cell, not a paragraph border)
_spacer(doc, 3)
hrtbl = doc.add_table(rows=1, cols=1)
set_col_widths(hrtbl, [CONTENT_W])
hrc = hrtbl.cell(0, 0)
_set_cell_margins(hrc, top=0, bottom=0, left=0, right=0)
_set_cell_border(hrc, bottom=(16, NAVY), top=None, left=None, right=None)
_cell_para(hrc)
_spacer(doc, 6)

# Stat strip — the headline numbers from §1, as data tiles not prose
stat_tile_row(doc, [
    ("32.7 GW", "Operating + queued, 60 mi"),
    ("15.5 MI", "To GW Ranch (7.65 GW)"),
    ("19.3 MI", "To Longfellow (2 GW)"),
    ("47,418 AF/YR", "Water permitted, adjacent"),
])
_spacer(doc, 4)

# Positioning strip — tightened §1, one short paragraph, not the page's content
pp = doc.add_paragraph()
pp.paragraph_format.space_before = Pt(6)
pp.paragraph_format.space_after = Pt(4)
_run(pp,
     "1,300 contiguous acres on I-10 in Pecos County, positioned inside an already-forming "
     "power and data-center corridor: two hyperscale-scale projects (7.65 GW and 2 GW) sit on "
     "the same north–south line through the property at 15.5 and 19.3 miles, backed by "
     "transmission, water, and gas positions that are permitted or contract-ready, not proposed. "
     "The tables below carry the data-room figures; narrative framing is in the companion "
     "Executive Brief.",
     size=8.6, color=GRAPHITE)

# Two-column: gravity diagram (left) | Property facts (right)
P1_LEFT_W = 3.15
P1_RIGHT_W = CONTENT_W - 3.15 - 0.25
left_cell, right_cell = two_col(doc, left_w=P1_LEFT_W, right_w=P1_RIGHT_W)

section_tag_and_title(left_cell, "FIG 2.6", "Power-Gravity Diagram", width=P1_LEFT_W)
add_picture_fit(left_cell, EXHIBIT_DIR / "exhibit_power_gravity_light.png", 3.0)
caption(left_cell, "Concentric operating + ERCOT-queue capacity, Caramba North at center. "
                    "GW Ranch ≈N (19°); Longfellow ≈S (188°).")

section_tag_and_title(right_cell, "SEC 2.1", "The Property", width=P1_RIGHT_W)
insight_callout(right_cell,
    "As-of-right industrial land inside the fastest-growing load pocket in ERCOT — "
    "not a rezoning story.", width=P1_RIGHT_W)
fact_table(right_cell, [
    ["PROPERTY FACT", "VALUE"],
    ["Max contiguous acreage", "1,300 ac"],
    ["Location", "N side of I-10, Pecos Co."],
    ["Distance to Fort Stockton", "≈ 5 mi"],
    ["Zoning", "None — as-of-right"],
    ["ERCOT weather zone", "Far West (highest-growth)"],
], col_widths=[2.55, 1.4], aligns=["left", "right"], bold_first_col=False)

_spacer(doc, 8)

section_tag_and_title(doc, "SEC 2.2", "Transmission")
insight_callout(doc,
    "Fifteen miles from the delivery point of all three approved 765 kV Permian import lines "
    "— the transmission decision is already made, upstream of this site.")
fact_table(doc, [
    ["TRANSMISSION FACT", "VALUE"],
    ["Distance to Solstice Substation (AEP/CPS Energy)", "15 mi"],
    ["Role of Solstice", "Western terminus, 3 PUCT-approved 765 kV Permian import paths"],
    ["Approval", "PUCT Docket No. 55718, approved Apr 24, 2025"],
    ["Local substations within 10 mi", "6 (138/69 kV and 69 kV)"],
    ["ERCOT-wide TPIT pipeline (planned, not built)", "141 substation + 133 line upgrades"],
], col_widths=[3.55, 3.85], aligns=["left", "left"])

_spacer(doc, 10)

# ---------------------------------------------------------------- PAGE 2 ---
section_tag_and_title(doc, "SEC 2.3", "Regional Power Cluster — Pecos County")
insight_callout(doc,
    "12 GW already queued in this county alone — before counting the two hyperscale "
    "campuses profiled in §2.6.")
fact_table(doc, [
    ["CAPACITY", "OPERATING", "ERCOT QUEUE"],
    ["Pecos County — Solar", "2,178 MW / 13 proj.", "—"],
    ["Pecos County — Wind", "542 MW / 5 proj.", "—"],
    ["Pecos County — BESS", "505 MW / 6 proj.", "—"],
    ["Pecos County — Gas", "1 MW / 1 proj.", "—"],
    ["Pecos County — TOTAL", "3,226 MW", "12,039 MW / 39 proj."],
    ["Adjacent 6 counties (Reeves, Crane, Ward, Upton, Ector, Crockett)", "7,022 MW", "24,585 MW"],
    ["Within 20 mi of Caramba North specifically", "—", "3,973 MW / 13 proj."],
    ["Nearest operating storage", "St. Gall Energy Storage I — 1.9 mi, 103 MW BESS", "—"],
], col_widths=[3.9, 1.8, 1.7], aligns=["left", "right", "right"], bold_first_col=False)

_spacer(doc, 8)

section_tag_and_title(doc, "SEC 2.6", "Ring Analysis — Region-Wide Operating + Queued Capacity")
insight_callout(doc,
    "32.7 GW sits within 60 miles of Caramba North — the property is inside the radius, "
    "not adjacent to it.")
ring_rows = [["RADIUS", "OPERATING + QUEUE"]]
for r in RING:
    ring_rows.append([f"≤ {r['radius_mi']} mi", f"{r['total_gw']:.1f} GW"])
fact_table(doc, ring_rows, col_widths=[3.9, 3.5], aligns=["left", "right"])
small_print(doc, "Region-wide EIA-860 operating + ERCOT-queue capacity within each radius of the "
                  "Caramba North tract boundary — not county-bounded. See Fig. 2.6.")

_spacer(doc, 6)

section_tag_and_title(doc, "SEC 2.6", "Project-Maturity Framing")
mat = MATURITY
fact_table(doc, [
    ["STATUS", "SHARE OF COMBINED ANCHOR MW"],
    ["Under construction (GW Ranch)", f"{mat['under_construction']['pct_of_local_mw']:.1f}%"],
    ["Planned / phase-1 (Longfellow)", f"{mat['seeking_tenant']['pct_of_local_mw']:.1f}%"],
], col_widths=[3.9, 3.5], aligns=["left", "right"])
small_print(doc, "The regional pipeline anchored by the two profiled campuses (feature profiles "
                  "below) is majority-built, not majority-speculative.")

_spacer(doc, 10)

# ---------------------------------------------------------------- PAGE 3 ---
P3_COL_W = (CONTENT_W - 0.3) / 2
left_cell, right_cell = two_col(doc, left_w=P3_COL_W, right_w=P3_COL_W, gutter=0.3)

section_tag_and_title(left_cell, "SEC 2.4", "Water", width=P3_COL_W)
insight_callout(left_cell,
    "Two-thirds of the district's industrial water rights are already permitted to this "
    "position — the water conversation is closed, not open.", width=P3_COL_W)
fact_table(left_cell, [
    ["WATER FACT", "VALUE"],
    ["Permitted volume", "47,418 AF/yr"],
    ["Permitted volume", "42.3 MGD"],
    ["Share of Middle Pecos GCD industrial rights", "≈ 2/3"],
    ["Source aquifer", "Edwards-Trinity (Plateau)"],
    ["Recharge record", "Held through 1950s drought of record"],
], col_widths=[2.2, 1.35], aligns=["left", "right"])

section_tag_and_title(right_cell, "SEC 2.5", "Natural Gas", width=P3_COL_W)
insight_callout(right_cell,
    "A signable 15-year gas quote at Waha basis — the same structural discount now drawing "
    "behind-the-meter generation to this corridor.", width=P3_COL_W)
fact_table(right_cell, [
    ["GAS FACT", "VALUE"],
    ["Distance to Waha hub", "20 mi"],
    ["Indicative supply quote", "200,000 MMBtu/d"],
    ["Term", "15 yr, Waha-index"],
    ["CIAC", "$15–25M"],
    ["Lead time", "9–15 mo"],
], col_widths=[2.15, 1.4], aligns=["left", "right"])

_spacer(doc, 8)

# Macro context box — 1 of the 2 permitted uses in the document
mtbl = doc.add_table(rows=1, cols=1)
set_col_widths(mtbl, [CONTENT_W])
mc = mtbl.cell(0, 0)
_set_cell_bg(mc, "F4EFE3")
_set_cell_border(mc, top=None, bottom=None, right=None, left=(18, "B7862E"))
_set_cell_margins(mc, top=90, bottom=90, left=140, right=140)
mp = _cell_para(mc)
_run(mp, "MACRO CONTEXT — ", size=7.6, bold=True, color="7A5A19", caps=True, spacing=6)
_run(mp,
     "ERCOT's large-load interconnection queue grew from 63 GW (end 2024) to ~474 GW of pending "
     "requests by Aug 2026 (~90% data-center-driven) — large enough that the Governor "
     "directed an audit of all ERCOT-queue data centers and paused the “Batch Zero” "
     "large-load review (public reporting, Aug 2026). The 32.7 GW within 60 mi of Caramba North "
     "sits inside that same queue.",
     size=7.6, color="4A3A12")
_row_cant_split(mtbl.rows[0])

doc.add_page_break()

# ---------------------------------------------------------------- PAGE 4 ---
section_tag_and_title(doc, "SEC 2.6", "Regional Data-Center / Power Pipeline — Feature Profiles")

# GW Ranch
section_tag_and_title(doc, "2.6.A", "GW Ranch (Amazon)", space_before=4)
insight_callout(doc,
    "The largest air permit issued in the US this year sits fifteen miles up the same highway "
    "corridor — under construction, not announced.", space_after=4)
gl, gr = two_col(doc, left_w=4.35, right_w=CONTENT_W - 4.35 - 0.25)
fact_table(gl, [
    ["GW RANCH FACT", "VALUE"],
    ["Site size", "8,000 ac, Pecos County"],
    ["Distance to Caramba North (edge-to-edge)", "≈ 15.5 mi"],
    ["Ownership", "Amazon (disclosed Aug 2026)"],
    ["Generation permit", "7.65 GW TCEQ air permit (largest in US, issued Jan/Feb 2026)"],
    ["Turbines", "35 gas turbines"],
    ["Storage / solar", "1.8 GW battery; up to 750 MW solar"],
    ["Data-center buildings", "3 × 189,000 sf (Gensler); target Dec 2026"],
    ["Estimated investment", "≈ $12B total project"],
    ["Status", "Under construction"],
], col_widths=[2.7, 1.45], aligns=["left", "right"], feature=True)
add_picture_fit(gr, GWRANCH_IMG, CONTENT_W - 4.35 - 0.25)
caption(gr, "GW Ranch — ≈15.5 mi from Caramba North.")
small_print(gr, "TCEQ air permit is a generation permit, not an ERCOT queue position; no "
                "disclosed ERCOT filing — site is off-grid initially.", size=6.4)

_spacer(doc, 6)

# Longfellow
section_tag_and_title(doc, "2.6.B", "Longfellow", space_before=4)
insight_callout(doc,
    "A second phased gas-generation campus twenty miles south — the corridor's demand for "
    "on-site power isn't one project deep.", space_after=4)
ll, lr = two_col(doc, left_w=4.35, right_w=CONTENT_W - 4.35 - 0.25)
fact_table(ll, [
    ["LONGFELLOW FACT", "VALUE"],
    ["Site size", "568 ac, Pecos County"],
    ["Distance to Caramba North (edge-to-edge)", "≈ 19.3 mi"],
    ["On-site generation plan", "Aero-derivative gas turbines, SCR + carbon-capture capable"],
    ["Cooling", "Closed-loop, permitted non-potable groundwater"],
    ["Original scope (Oct 2025 announcement)", "2 GW, 8 phases (250 MW/phase)"],
    ["Status", "Phase-1 site work underway; generation build planned in phases"],
    ["ERCOT queue / TCEQ air permit on record", "None found as of Aug 2026"],
], col_widths=[2.7, 1.45], aligns=["left", "right"], feature=True)
add_picture_fit(lr, LONGFELLOW_IMG, CONTENT_W - 4.35 - 0.25)
caption(lr, "Longfellow — ≈19.3 mi from Caramba North.")
small_print(lr, "Longfellow's own public materials describe the site as “more than 25 mi "
                "outside Fort Stockton” — consistent with the longer edge-to-edge figure.",
                size=6.4)

doc.add_page_break()

# ---------------------------------------------------------------- PAGE 5 ---
section_tag_and_title(doc, "SEC 2.8", "Subsurface — New-Drill Activity, Pecos County")
insight_callout(doc,
    "Pecos County has the lowest new-drilling count of seven comparable Permian counties since "
    "2020 — a 90%-below-peer-average level of activity, not merely “quiet.”")
fact_table(doc, [
    ["SUBSURFACE FACT", "VALUE"],
    ["New-drill wellbore events since 2020, Pecos County", "115 of 1,140 total events (10%)"],
    ["Share of Pecos events that are workovers/reworks", "90%"],
    ["New-drill wells within 2 mi of tract, since 2020", "0"],
    ["New-drill wells within 5 mi of tract, since 2020", "0"],
    ["New-drill wells within 10 mi of tract, since 2020", "1 (9.37 mi)"],
    ["New-drill wells beyond 10 mi, since 2020", "114 (median 19.9 mi, mean 20.9 mi)"],
    ["Non-plugged wellbores that are marginal/EOL, ≤ 10 mi", "83% (vs. 60% ≤ 2 mi, 62% ≤ 5 mi)"],
], col_widths=[3.85, 3.55], aligns=["left", "left"], bold_first_col=False)

_spacer(doc, 5)

section_tag_and_title(doc, "SEC 2.8", "Peer-County Comparison — New-Drill Count Since 2020", space_before=2)
fact_table(doc, [
    ["COUNTY", "NEW-DRILL COUNT"],
    ["Pecos (subject county)", "115"],
    ["Reagan", "668"],
    ["Howard", "990"],
    ["Reeves", "1,053"],
    ["Loving", "1,121"],
    ["Midland", "1,569"],
    ["Martin", "1,685"],
    ["Peer average (6 counties, excl. Pecos)", "1,181"],
], col_widths=[3.9, 3.5], aligns=["left", "right"], bold_first_col=False)

_spacer(doc, 8)

section_tag_and_title(doc, "SEC 2.7", "The Diligence Platform", space_before=2)
insight_callout(doc,
    "Every figure in this document is independently re-derivable from a cited public source "
    "— this isn't a broker's summary.", space_after=4)
fact_table(doc, [
    ["PLATFORM FACT", "VALUE"],
    ["Source datasets", "ERCOT GIS Report/TPIT, PUCT, EIA-860, TCEQ, RRC, FracFocus, Middle Pecos GCD, HIFLD, USGS, BTS, Census TIGER"],
    ["Refresh cadence", "RRC weekly; ERCOT queue/TPIT monthly; EIA/USGS/OSM annually"],
    ["Build", "Static, versioned; deployed bundle byte-verified on release; access logged"],
    ["Access", "lrp-tx-gis.netlify.app (credentials issued to deal team separately)"],
], col_widths=[1.85, 5.55], aligns=["left", "left"])

_spacer(doc, 10)

# Notices / disclaimer
ntbl = doc.add_table(rows=1, cols=1)
set_col_widths(ntbl, [CONTENT_W])
nc = ntbl.cell(0, 0)
_set_cell_bg(nc, "EFF1F2")
_set_cell_border(nc, top=(4, LINE), bottom=(4, LINE), left=(4, LINE), right=(4, LINE))
_set_cell_margins(nc, top=90, bottom=90, left=130, right=130)
np_ = _cell_para(nc)
_run(np_, "NOTICES", size=6.8, bold=True, color=MUTED, caps=True, spacing=8)
np2 = nc.add_paragraph()
np2.paragraph_format.space_before = Pt(2)
_run(np2,
     "Confidential offering memorandum prepared for a limited number of prospective "
     "counterparties under NDA. Not an offer to sell or a solicitation of securities. "
     "Information is preliminary and indicative, from sources believed reliable. Public data is "
     "drawn from the sources listed under §2.7 above; third-party transaction news is sourced "
     "to public reporting cited in the companion source register.",
     size=6.6, color=MUTED)
_row_cant_split(ntbl.rows[0])

small_print(doc,
    "Distance methodology: distances to GW Ranch and Longfellow are measured edge-to-edge, from "
    "the nearest point on the Caramba North tract boundary to each site's disclosed location, "
    "not centroid-to-centroid — consistently shorter given the tract's own spatial extent. "
    "GW Ranch: 15.5 mi (vs. 17.3 mi centroid). Longfellow: 19.3 mi (vs. 19.7 mi centroid); "
    "Longfellow's own public site states its location as more than 25 mi outside Fort Stockton, "
    "consistent with the longer figure — this distance should not be represented as shorter.",
    size=6.4, space_before=6)

# python-docx's default template ships a <w:zoom> with no required w:percent
# attribute, which fails strict OOXML validation — patch it before saving.
zoom_el = doc.settings.element.find(qn("w:zoom"))
if zoom_el is not None and zoom_el.get(qn("w:percent")) is None:
    zoom_el.set(qn("w:percent"), "100")

doc.save(OUT)
print(f"wrote {OUT}")
